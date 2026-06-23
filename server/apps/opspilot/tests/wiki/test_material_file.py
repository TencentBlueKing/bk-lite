"""文件资料解析(.txt/.md)测试。MinIO 读取经 _read_file 注入,避免依赖对象存储。"""

import pytest

from apps.opspilot.services.wiki import material_service


def _kb():
    from apps.opspilot.models import WikiKnowledgeBase

    return WikiKnowledgeBase.objects.create(name="kb", team=[1])


def _file_material(kb, name):
    from apps.opspilot.models import Material

    return Material.objects.create(knowledge_base=kb, name=name, material_type="file")


@pytest.mark.django_db
def test_extract_markdown_file(monkeypatch):
    kb = _kb()
    mat = _file_material(kb, "guide.md")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("guide.md", b"# \xe6\xa0\x87\xe9\xa2\x98\nbody"))
    text = material_service.extract_text(mat)
    assert "标题" in text and "body" in text


@pytest.mark.django_db
def test_extract_txt_file(monkeypatch):
    kb = _kb()
    mat = _file_material(kb, "notes.txt")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("notes.txt", b"plain text content"))
    assert material_service.extract_text(mat) == "plain text content"


@pytest.mark.django_db
def test_unsupported_extension_returns_empty(monkeypatch):
    kb = _kb()
    mat = _file_material(kb, "scan.pdf")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("scan.pdf", b"%PDF-1.4 ..."))
    assert material_service.extract_text(mat) == ""


@pytest.mark.django_db
def test_ingest_markdown_file_sets_done(monkeypatch):
    kb = _kb()
    mat = _file_material(kb, "doc.md")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("doc.md", b"hello wiki"))
    out = material_service.ingest_material(mat)
    assert out.status == "done"
    assert out.ai_summary == "hello wiki"  # 无模型回退为截断正文
    assert out.content_hash


def _xlsx_bytes():
    import io

    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["host", "status"])
    ws.append(["web01", "running"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@pytest.mark.django_db
def test_extract_xlsx_file(monkeypatch):
    kb = _kb()
    mat = _file_material(kb, "data.xlsx")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("data.xlsx", _xlsx_bytes()))
    text = material_service.extract_text(mat)
    assert "web01" in text and "running" in text


@pytest.mark.django_db
def test_ocr_dispatch_uses_provider_and_loader(monkeypatch):
    from langchain_core.documents import Document

    kb = _kb()
    mat = _file_material(kb, "scan.pdf")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("scan.pdf", b"%PDF-1.4 fake"))
    monkeypatch.setattr(material_service, "_build_ocr", lambda m: object())  # 模拟已配置 OCR

    class FakeLoader:
        def __init__(self, path, ocr, mode):
            self.ocr = ocr

        def load(self):
            return [Document("ocr 提取的文本")]

    monkeypatch.setattr(material_service, "_ocr_loader_class", lambda ext: FakeLoader)
    assert material_service.extract_text(mat) == "ocr 提取的文本"


@pytest.mark.django_db
def test_image_without_provider_returns_empty(monkeypatch):
    # 图片纯图像内容,必须 OCR;无 OCRProvider → 空串
    kb = _kb()
    mat = _file_material(kb, "scan.png")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("scan.png", b"\x89PNG\r\n\x1a\n fake"))
    monkeypatch.setattr(material_service, "_build_ocr", lambda m: None)
    assert material_service.extract_text(mat) == ""


def _docx_bytes():
    import io

    from docx import Document as Docx

    d = Docx()
    d.add_paragraph("运维手册标题")
    d.add_paragraph("重启服务的步骤说明")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@pytest.mark.django_db
def test_docx_native_text_without_ocr(monkeypatch):
    # 文档型 .docx:原生抽取文本,无需 OCR 服务(OCR 仅增强内嵌图片)
    kb = _kb()
    mat = _file_material(kb, "manual.docx")
    monkeypatch.setattr(material_service, "_read_file", lambda m: ("manual.docx", _docx_bytes()))
    monkeypatch.setattr(material_service, "_build_ocr", lambda m: None)  # 无 OCRProvider
    text = material_service.extract_text(mat)
    assert "运维手册标题" in text and "重启服务的步骤说明" in text


@pytest.mark.django_db
def test_extract_web_text(monkeypatch):
    from apps.opspilot.models import Material

    kb = _kb()
    mat = Material.objects.create(knowledge_base=kb, name="site", material_type="web", url="http://example.com")
    monkeypatch.setattr(
        material_service,
        "_fetch_url",
        lambda url: "<html><body><h1>标题</h1><p>正文内容</p><script>bad()</script></body></html>",
    )
    text = material_service.extract_text(mat)
    assert "标题" in text and "正文内容" in text and "bad" not in text


@pytest.mark.django_db
def test_extract_web_text_fetch_failure_returns_empty(monkeypatch):
    from apps.opspilot.models import Material

    kb = _kb()
    mat = Material.objects.create(knowledge_base=kb, name="site", material_type="web", url="http://example.com")

    def _boom(url):
        raise RuntimeError("network down")

    monkeypatch.setattr(material_service, "_fetch_url", _boom)
    assert material_service.extract_text(mat) == ""


@pytest.mark.integration
@pytest.mark.django_db
def test_material_file_upload_endpoint(api_client):
    """前后联动:前端 multipart 文档上传 → POST /material/ → 文件落 MinIO → 资料创建。"""
    from django.core.files.uploadedfile import SimpleUploadedFile

    kb = _kb()
    resp = api_client.post(
        "/api/v1/opspilot/wiki_mgmt/material/",
        {
            "knowledge_base": kb.id,
            "name": "guide.md",
            "material_type": "file",
            "file": SimpleUploadedFile("guide.md", b"# \xe6\x8c\x87\xe5\x8d\x97\nbody"),
        },
        format="multipart",
    )
    assert resp.status_code == 201
    data = resp.json()["data"]
    assert data["material_type"] == "file" and data["id"]


@pytest.mark.integration
@pytest.mark.django_db
def test_ingest_xlsx_via_real_minio():
    """真实 MinIO 往返:上传 xlsx 到对象存储,再读回解析,验证 _read_file 端到端。"""
    from django.core.files.uploadedfile import SimpleUploadedFile

    from apps.opspilot.models import Material

    kb = _kb()
    mat = Material.objects.create(
        knowledge_base=kb,
        name="real.xlsx",
        material_type="file",
        file=SimpleUploadedFile("real.xlsx", _xlsx_bytes()),
    )
    out = material_service.ingest_material(mat)
    assert out.status == "done"
    assert "web01" in out.ai_summary
