"""K8S 配置检查报告 .docx 生成器（参考 K8S集群配置检查报告.docx 模板样式）

报告结构：
1. 封面标题 + 集群信息
2. 总体概览（统计表格）
3. 按严重度分组的问题详情（编号 + 表格 + 风险/修复文字）
4. 修复优先级建议（P0-P3 分组编号列表，不列具体 workload）
5. 页脚

不包含修复命令（命令通过前端事件单独展示）。
"""

import base64
import uuid
from collections import OrderedDict
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn


# 严重度配置
_SEVERITY_CONFIG = OrderedDict([
    ("critical", {"label": "严重问题 (Critical)", "color": RGBColor(0xCC, 0x00, 0x00), "desc": "需立即修复", "priority": "P0 — 立即修复（安全风险）"}),
    ("high", {"label": "高危问题 (High)", "color": RGBColor(0xE6, 0x5C, 0x00), "desc": "建议本周内修复", "priority": "P1 — 本周完成（生产稳定性）"}),
    ("warning", {"label": "中等问题 (Medium)", "color": RGBColor(0xCC, 0x99, 0x00), "desc": "建议两周内修复", "priority": "P2 — 两周内完成（安全加固）"}),
    ("info", {"label": "低级告警 (Low)", "color": RGBColor(0x33, 0x66, 0xCC), "desc": "持续改进", "priority": "P3 — 持续改进"}),
])

_TABLE_HEADER_BG = "1F4E79"
_TABLE_HEADER_COLOR = RGBColor(0xFF, 0xFF, 0xFF)


def generate_k8s_report_docx(
    report_data: Dict[str, Any],
    commands_text: str = "",
) -> bytes:
    """生成 K8S 配置检查报告 .docx 文件（不包含修复命令）

    report_data 结构：
      - cluster_name: str
      - raw_items: List[Dict] - 每个 item 包含:
          namespace, target_name, target_type, severity, summary, category
    """
    doc = Document()

    # 设置默认字体为微软雅黑
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style_rpr = style.element.get_or_add_rPr()
    style_rpr_fonts = style_rpr.find(qn("w:rFonts"))
    if style_rpr_fonts is None:
        from docx.oxml import OxmlElement
        style_rpr_fonts = OxmlElement("w:rFonts")
        style_rpr.append(style_rpr_fonts)
    style_rpr_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    cluster_name = report_data.get("cluster_name", "未知集群")
    # 支持 raw_items（原始逐条数据）和 items（旧的聚合数据）
    items = report_data.get("raw_items", report_data.get("items", []))
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M GMT+8")

    # 收集涉及的命名空间
    namespaces = sorted(set(
        item.get("namespace", "") for item in items if item.get("namespace")
    ))
    ns_text = ", ".join(namespaces) if namespaces else "default"

    # === 1. 封面标题 ===
    _add_styled_para(doc, "Kubernetes", size=36, bold=True, color=RGBColor(0x1F, 0x4E, 0x79), align="center")
    _add_styled_para(doc, "集群工作负载配置检查报告", size=28, bold=True, color=RGBColor(0x1F, 0x4E, 0x79), align="center")
    _add_styled_para(doc, "—— 安全与最佳实践审计 ——", size=14, color=RGBColor(0x66, 0x66, 0x66), align="center")
    doc.add_paragraph()

    # 集群信息
    _add_styled_para(doc, "集群信息", size=12, bold=True, color=RGBColor(0x40, 0x40, 0x40))
    _add_styled_para(doc, f"集群：{cluster_name}", size=11, color=RGBColor(0x66, 0x66, 0x66))
    _add_styled_para(doc, f"检查时间：{generated_at}", size=11, color=RGBColor(0x66, 0x66, 0x66))
    _add_styled_para(doc, f"检查范围：{ns_text}", size=11, color=RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()

    # === 2. 总体概览 ===
    doc.add_heading("总体概览", level=1)
    _add_styled_para(doc, "本报告对集群内工作负载进行了安全与配置最佳实践审计。", size=11)

    severity_counts = _count_severities(items)
    total_issues = sum(severity_counts.values())
    workload_count = len(set(
        f"{i.get('namespace','')}/{i.get('target_name', i.get('workload_name',''))}"
        for i in items
    ))

    overview_table = doc.add_table(rows=1, cols=3)
    overview_table.style = "Table Grid"
    _set_header_row(overview_table.rows[0], ["指标", "数量", "说明"])

    _add_table_row(overview_table, ["涉及工作负载", str(workload_count), f"分布在 {len(namespaces)} 个命名空间"])
    for sev_key in ["critical", "high", "warning", "info"]:
        count = severity_counts.get(sev_key, 0)
        if count == 0:
            continue
        cfg = _SEVERITY_CONFIG[sev_key]
        _add_table_row(overview_table, [cfg["label"], str(count), cfg["desc"]])
    _add_table_row(overview_table, ["问题总数", str(total_issues), ""])

    doc.add_paragraph()

    # === 3. 按严重度分组的问题详情 ===
    items_by_severity: Dict[str, List] = {}
    for item in items:
        sev = item.get("severity", "info")
        items_by_severity.setdefault(sev, []).append(item)

    issue_idx = 1
    for sev_key in ["critical", "high", "warning", "info"]:
        sev_items = items_by_severity.get(sev_key, [])
        if not sev_items:
            continue

        cfg = _SEVERITY_CONFIG[sev_key]
        doc.add_heading(cfg["label"], level=1)

        sev_descs = {
            "critical": "以下问题存在直接的安全风险，必须立即修复。",
            "high": "以下问题影响生产稳定性，建议尽快修复。",
            "warning": "以下问题不符合最佳实践，建议两周内修复。",
            "info": "以下为低优先级改进建议，可持续改进。",
        }
        _add_styled_para(doc, sev_descs.get(sev_key, ""), size=11, color=cfg["color"])

        # 按问题类型（summary）分组 — 每种问题一个编号标题 + 表格列出所有受影响的 workload
        issues_by_type: Dict[str, List] = {}
        for item in sev_items:
            summary = item.get("summary", "未知问题")
            issues_by_type.setdefault(summary, []).append(item)

        for issue_type, type_items in issues_by_type.items():
            doc.add_heading(f"{issue_idx}. {issue_type}", level=2)

            if len(type_items) == 1:
                # 单个工作负载 — 用描述文本
                it = type_items[0]
                target = it.get("target_name", it.get("workload_name", ""))
                namespace = it.get("namespace", "")
                _add_styled_para(doc, f"{namespace}/{target} 存在此问题。", size=11)
            else:
                # 多个工作负载 — 用表格列出每一个
                tbl = doc.add_table(rows=1, cols=3)
                tbl.style = "Table Grid"
                _set_header_row(tbl.rows[0], ["工作负载", "命名空间", "问题详情"])

                for it in type_items:
                    target = it.get("target_name", it.get("workload_name", ""))
                    namespace = it.get("namespace", "")
                    detail = it.get("summary", "")
                    _add_table_row(tbl, [target, namespace, detail])

            # 风险 + 修复
            risk_text = _get_risk_description(issue_type)
            fix_text = _get_fix_description(issue_type)

            para = doc.add_paragraph()
            run_label = para.add_run("风险：")
            run_label.bold = True
            run_label.font.size = Pt(11)
            run_label.font.name = "Microsoft YaHei"
            run_label.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run_content = para.add_run(risk_text)
            run_content.font.size = Pt(11)
            run_content.font.name = "Microsoft YaHei"
            run_content.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

            para = doc.add_paragraph()
            run_label = para.add_run("修复：")
            run_label.bold = True
            run_label.font.size = Pt(11)
            run_label.font.name = "Microsoft YaHei"
            run_label.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run_content = para.add_run(fix_text)
            run_content.font.size = Pt(11)
            run_content.font.name = "Microsoft YaHei"
            run_content.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

            issue_idx += 1

    # === 4. 修复优先级建议 ===
    doc.add_heading("修复优先级建议", level=1)

    priority_idx = 1
    for sev_key in ["critical", "high", "warning", "info"]:
        sev_items = items_by_severity.get(sev_key, [])
        if not sev_items:
            continue

        cfg = _SEVERITY_CONFIG[sev_key]
        doc.add_heading(cfg["priority"], level=2)

        # 按问题类型去重，只列问题类型和数量
        issues_by_type_p: Dict[str, int] = {}
        for item in sev_items:
            summary = item.get("summary", "")
            issues_by_type_p[summary] = issues_by_type_p.get(summary, 0) + 1

        for issue_type, count in issues_by_type_p.items():
            _add_styled_para(doc, f"{priority_idx}. {issue_type}（{count} 个工作负载）", size=11)
            priority_idx += 1

    # === 5. 页脚 ===
    doc.add_paragraph()
    _add_styled_para(doc, "— 报告结束 —", size=10, color=RGBColor(0x99, 0x99, 0x99), align="center")
    _add_styled_para(
        doc, f"由 WeOpsX 自动生成 | {generated_at}",
        size=9, color=RGBColor(0xAA, 0xAA, 0xAA), align="center"
    )

    # === 导出 ===
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_report_download_event(
    report_data: Dict[str, Any],
    commands_text: str = "",
) -> Dict[str, str]:
    """生成报告并返回可直接用于 dispatch 的事件数据"""
    try:
        docx_bytes = generate_k8s_report_docx(report_data)
    except Exception as e:
        import logging
        logging.getLogger(__name__).error("docx 报告生成失败: %s", str(e), exc_info=True)
        return {
            "error": f"报告生成失败: {str(e)}",
        }

    content_b64 = base64.b64encode(docx_bytes).decode("utf-8")

    cluster_name = report_data.get("cluster_name", "集群")
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"K8S配置检查报告_{cluster_name}_{date_str}.docx"

    return {
        "download_id": str(uuid.uuid4())[:8],
        "filename": filename,
        "content_base64": content_b64,
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }


# === 辅助函数 ===


def _add_styled_para(doc, text: str, size: float = 11, bold: bool = False,
                     color: RGBColor = None, align: str = None):
    """添加样式化段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para


def _set_header_row(row, texts: List[str]):
    """设置表格表头行（深蓝背景白字）"""
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Microsoft YaHei"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.bold = True
        run.font.color.rgb = _TABLE_HEADER_COLOR
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {})
        shd.set(qn("w:fill"), _TABLE_HEADER_BG)
        shd.set(qn("w:val"), "clear")
        shading.append(shd)


def _add_table_row(table, texts: List[str]):
    """添加表格数据行"""
    row = table.add_row()
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "Microsoft YaHei"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _count_severities(items: List[Dict]) -> Dict[str, int]:
    """统计各严重度数量"""
    counts: Dict[str, int] = {}
    for item in items:
        sev = item.get("severity", "info")
        counts[sev] = counts.get(sev, 0) + 1
    return counts


def _get_risk_description(issue_type: str) -> str:
    """根据问题类型返回风险描述"""
    if "资源限制" in issue_type or "资源请求" in issue_type:
        return "无资源限制的容器可消耗节点所有资源，导致其他 Pod OOM 或 CPU 饥饿，影响集群稳定性。"
    if "存活探针" in issue_type or "健康探针" in issue_type or "liveness" in issue_type.lower():
        return "无存活探针时 Kubernetes 无法自动检测和重启不健康的容器，故障容器将持续运行。"
    if "就绪探针" in issue_type or "readiness" in issue_type.lower():
        return "无就绪探针时 Service 可能将流量路由到未准备好的 Pod，导致请求失败。"
    if "探针" in issue_type:
        return "缺少健康检查探针，Kubernetes 无法自动检测容器故障并执行自愈操作。"
    if "root" in issue_type or "安全上下文" in issue_type or "非 root" in issue_type.lower():
        return "容器以 root 用户运行，容器逃逸后攻击者将获得宿主机 root 权限，安全风险极高。"
    if "latest" in issue_type or "镜像标签" in issue_type:
        return "latest 标签不可溯源，每次拉取可能获得不同版本，导致不可预测行为和回滚困难。"
    if "单副本" in issue_type or "副本" in issue_type:
        return "单副本部署存在单点故障风险，节点异常时服务将完全中断，无法保障高可用。"
    if "特权" in issue_type or "privileged" in issue_type.lower():
        return "特权容器拥有宿主机全部 Linux capabilities，容器逃逸后等同于 root 访问整个节点。"
    if "hostNetwork" in issue_type or "主机命名空间" in issue_type or "hostPID" in issue_type:
        return "共享宿主机网络/进程/IPC 命名空间会绕过网络和进程隔离，增大攻击面。"
    if "密码" in issue_type or "明文" in issue_type or "Secret" in issue_type:
        return "密码暴露在 Git 历史、kubectl describe 输出中，任何有 namespace 读权限的用户均可看到。"
    if "NetworkPolicy" in issue_type or "网络隔离" in issue_type:
        return "所有 Pod 之间可自由通信，一旦某个容器被入侵，攻击者可横向移动到所有命名空间。"
    if "ServiceAccount" in issue_type:
        return "使用默认 ServiceAccount 违反最小权限原则，可能被利用进行集群内横向攻击。"
    return "当前配置不符合 Kubernetes 最佳实践，可能影响集群安全性和稳定性。"


def _get_fix_description(issue_type: str) -> str:
    """根据问题类型返回修复建议"""
    if "资源限制" in issue_type or "资源请求" in issue_type:
        return "为所有容器设置 resources.requests 和 resources.limits，建议 CPU 100m-500m，内存 128Mi-256Mi。"
    if "存活探针" in issue_type or "liveness" in issue_type.lower():
        return "添加 livenessProbe 配置（建议 httpGet 方式），设置合理的 initialDelaySeconds 和 periodSeconds。"
    if "就绪探针" in issue_type or "readiness" in issue_type.lower():
        return "添加 readinessProbe 配置，确保 Pod 准备好接收流量后才加入 Service 端点。"
    if "探针" in issue_type:
        return "为容器添加 livenessProbe 和 readinessProbe，建议使用 httpGet 或 tcpSocket 检测方式。"
    if "root" in issue_type or "安全上下文" in issue_type or "非 root" in issue_type.lower():
        return "配置 securityContext.runAsNonRoot: true 和 runAsUser: 1000，禁止容器以 root 运行。"
    if "latest" in issue_type or "镜像标签" in issue_type:
        return "将所有 :latest 标签替换为固定版本号（如 :1.25.3），确保镜像版本可追溯。"
    if "单副本" in issue_type or "副本" in issue_type:
        return "将生产环境工作负载副本数增加至 2 或以上，配合 PodDisruptionBudget 保障高可用。"
    if "特权" in issue_type or "privileged" in issue_type.lower():
        return "移除 privileged: true，仅授予必要的 capabilities（如 NET_BIND_SERVICE）。"
    if "hostNetwork" in issue_type or "主机命名空间" in issue_type or "hostPID" in issue_type:
        return "移除 hostNetwork/hostPID/hostIPC 配置，使用 Service/Ingress 暴露服务端口。"
    if "密码" in issue_type or "明文" in issue_type or "Secret" in issue_type:
        return "使用 Secret 资源存储敏感信息，通过 secretKeyRef 引用，避免明文写入环境变量。"
    if "NetworkPolicy" in issue_type or "网络隔离" in issue_type:
        return "部署 default-deny NetworkPolicy，仅允许必要的 Pod 间通信。"
    if "ServiceAccount" in issue_type:
        return "为每个工作负载创建专用 ServiceAccount，配合 RBAC 最小权限策略。"
    return "参考 Kubernetes 安全基线调整配置。"
