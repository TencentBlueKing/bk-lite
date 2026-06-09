from __future__ import annotations

from datetime import timedelta
from io import BytesIO
from typing import Optional
from uuid import uuid4

from django.conf import settings
from django.core import signing
from django.core.files.base import ContentFile
from django.utils import timezone
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from apps.core.logger import opspilot_logger as logger
from apps.opspilot.models import FileKnowledge, WorkflowAttachmentAsset

# 工作流附件下载令牌的签名盐值与默认有效期(秒)。
# 使用签名 + 过期时间替代原先可被猜测且永不失效的明文 token。
WORKFLOW_ATTACHMENT_DOWNLOAD_SALT = "opspilot.workflow_attachment.download"


def workflow_attachment_download_max_age() -> int:
    """下载令牌最大有效期(秒)，默认 24 小时，可通过 settings 覆盖。"""
    return int(getattr(settings, "WORKFLOW_ATTACHMENT_DOWNLOAD_MAX_AGE", 24 * 60 * 60))


def build_signed_attachment_download_url(asset: WorkflowAttachmentAsset) -> str:
    """为工作流附件生成带签名且会过期的下载 URL。

    令牌绑定附件主键 id 与 execution_id，使用 TimestampSigner 机制(signing.dumps)
    携带时间戳，下载时按 max_age 校验过期并核对绑定关系，防止令牌被猜测或越权复用。
    URL 仍为 path 形式，保证聊天中的锚点链接无需额外请求头即可工作。
    """
    token = signing.dumps(
        {"aid": asset.id, "eid": asset.execution_id},
        salt=WORKFLOW_ATTACHMENT_DOWNLOAD_SALT,
    )
    return f"/api/v1/opspilot/bot_mgmt/workflow_attachment/download/{token}/"


def resolve_signed_attachment_token(download_token: str) -> Optional[WorkflowAttachmentAsset]:
    """校验签名下载令牌并返回绑定的附件，过期/非法/不匹配时抛出 BadSignature。"""
    payload = signing.loads(
        download_token,
        salt=WORKFLOW_ATTACHMENT_DOWNLOAD_SALT,
        max_age=workflow_attachment_download_max_age(),
    )
    if not isinstance(payload, dict):
        raise signing.BadSignature("Malformed download token payload")
    asset = WorkflowAttachmentAsset.objects.filter(id=payload.get("aid")).select_related("file_knowledge").first()
    # 校验签名内绑定的 execution_id，防止令牌被篡改后指向其它附件
    if not asset or asset.execution_id != payload.get("eid"):
        return None
    return asset

ATTACHMENT_FILE_TYPE_CONFIG = {
    "md": {
        "extension": "md",
        "mime_type": "text/markdown",
    },
    "pdf": {
        "extension": "pdf",
        "mime_type": "application/pdf",
    },
    "docx": {
        "extension": "docx",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    },
    "word": {
        "extension": "docx",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    },
}


def normalize_attachment_file_type(file_type: str, filename: str = "") -> tuple[str, str, str]:
    normalized_type = (file_type or "").strip().lower()
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if not normalized_type and extension:
        normalized_type = extension

    if normalized_type not in ATTACHMENT_FILE_TYPE_CONFIG:
        raise ValueError(f"不支持的附件类型: {file_type or extension or 'unknown'}")

    config = ATTACHMENT_FILE_TYPE_CONFIG[normalized_type]
    normalized_filename = filename or f"attachment.{config['extension']}"
    if "." not in normalized_filename:
        normalized_filename = f"{normalized_filename}.{config['extension']}"

    return normalized_type, normalized_filename, config["mime_type"]


def build_attachment_bytes(content: str, file_type: str, title: str = "") -> bytes:
    normalized_type, _, _ = normalize_attachment_file_type(file_type)
    safe_content = content or ""
    safe_title = title or "Attachment"

    if normalized_type == "md":
        return safe_content.encode("utf-8")
    if normalized_type == "pdf":
        return _build_pdf_bytes(safe_content, safe_title)
    if normalized_type in ("docx", "word"):
        return _build_docx_bytes(safe_content, safe_title)

    raise ValueError(f"不支持的附件类型: {file_type}")


def create_workflow_attachment_asset(
    *,
    execution_id: str,
    attachment_id: str,
    filename: str,
    content_bytes: bytes,
    mime_type: str,
    source_node_id: str = "",
    flow_id: str = "",
    created_by: str = "",
) -> WorkflowAttachmentAsset:
    if not execution_id:
        raise ValueError("execution_id 不能为空")
    if not attachment_id:
        raise ValueError("attachment_id 不能为空")

    file_knowledge = FileKnowledge.objects.create(file=ContentFile(content_bytes, name=filename))
    existing_asset = WorkflowAttachmentAsset.objects.filter(execution_id=execution_id, attachment_id=attachment_id).first()
    previous_file_knowledge = existing_asset.file_knowledge if existing_asset else None
    asset, _ = WorkflowAttachmentAsset.objects.update_or_create(
        execution_id=execution_id,
        attachment_id=attachment_id,
        defaults={
            "flow_id": flow_id,
            "source_node_id": source_node_id,
            "filename": filename,
            "mime_type": mime_type,
            "file_knowledge": file_knowledge,
            "created_by": str(created_by or ""),
        },
    )
    if previous_file_knowledge and previous_file_knowledge.id != asset.file_knowledge_id:
        previous_file_knowledge.delete()
    return asset


def build_workflow_attachment_id(*, execution_id: str, source_node_id: str = "", requested_attachment_id: str = "") -> str:
    normalized_attachment_id = str(requested_attachment_id or "").strip()
    if normalized_attachment_id:
        return normalized_attachment_id

    normalized_node_id = str(source_node_id or "").strip()
    if not execution_id or not normalized_node_id:
        return uuid4().hex[:12]

    existing_ids = set(
        WorkflowAttachmentAsset.objects.filter(execution_id=execution_id, source_node_id=normalized_node_id).values_list("attachment_id", flat=True)
    )
    if normalized_node_id not in existing_ids:
        return normalized_node_id

    suffix = 1
    while True:
        candidate = f"{normalized_node_id}__{suffix}"
        if candidate not in existing_ids:
            return candidate
        suffix += 1


def cleanup_expired_workflow_attachments(*, retention_days: int = 3) -> int:
    expire_before = timezone.now() - timedelta(days=retention_days)
    expired_assets = WorkflowAttachmentAsset.objects.filter(created_at__lt=expire_before).select_related("file_knowledge")
    deleted_count = 0

    for asset in expired_assets.iterator():
        if asset.file_knowledge_id:
            asset.file_knowledge.delete()
        else:
            asset.delete()
        deleted_count += 1

    return deleted_count


def _build_docx_bytes(content: str, title: str) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for line in content.splitlines() or [""]:
        document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf_bytes(content: str, title: str) -> bytes:
    font_name = _resolve_pdf_font()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=title,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "AttachmentTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=18,
        leading=24,
        spaceAfter=16,
    )
    body_style = ParagraphStyle(
        "AttachmentBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=16,
        spaceAfter=8,
    )
    story = [Paragraph(_escape_pdf_text(title), title_style), Spacer(1, 0.4 * cm)]
    for line in content.splitlines() or [""]:
        if line.strip():
            story.append(Paragraph(_escape_pdf_text(line), body_style))
        else:
            story.append(Spacer(1, 0.2 * cm))
    doc.build(story)
    return buffer.getvalue()


def _escape_pdf_text(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _resolve_pdf_font() -> str:
    font_candidates = [
        ("微软雅黑", "C:/Windows/Fonts/msyh.ttf"),
        ("微软雅黑", "C:/Windows/Fonts/msyh.ttc"),
        ("黑体", "C:/Windows/Fonts/simhei.ttf"),
        ("宋体", "C:/Windows/Fonts/simsun.ttc"),
    ]
    for font_name, font_path in font_candidates:
        try:
            pdfmetrics.registerFont(TTFont(font_name, font_path))
            return font_name
        except Exception:
            continue
    logger.warning("工作流附件 PDF 未找到可用中文字体，回退 Helvetica")
    return "Helvetica"
